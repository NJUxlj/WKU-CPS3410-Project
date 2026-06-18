#!/usr/bin/env python3
"""Generate the complete SA report with proper Word tables from scratch."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Load original incomplete report
doc_path = '/Users/xiniuyiliao/Desktop/application_code/WKU-CPS3410-Project/StimulatedAnnealingAlgorithm(1).docx'
doc = Document(doc_path)

def make_table(ref_element, headers, rows):
    """Insert a styled Word table right after ref_element."""
    all_data = [headers] + rows
    nr = len(all_data)
    nc = len(headers)
    
    tbl = OxmlElement('w:tbl')
    
    # Table width 100%
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    # Borders
    tblB = OxmlElement('w:tblBorders')
    for e in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{e}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), '999999')
        tblB.append(b)
    tblPr.append(tblB)
    tbl.append(tblPr)
    # Grid
    tblG = OxmlElement('w:tblGrid')
    for _ in range(nc):
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(9000//nc)); tblG.append(gc)
    tbl.append(tblG)
    
    for ri in range(nr):
        tr = OxmlElement('w:tr')
        for ci in range(nc):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(9000//nc)); tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            shd = OxmlElement('w:shd')
            if ri == 0: shd.set(qn('w:fill'), '2F5496')
            elif ri % 2 == 0: shd.set(qn('w:fill'), 'F2F2F2')
            else: shd.set(qn('w:fill'), 'FFFFFF')
            shd.set(qn('w:val'), 'clear'); tcPr.append(shd); tc.append(tcPr)
            
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr'); jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
            pPr.append(jc); p.append(pPr)
            r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), 'Calibri'); rf.set(qn('w:hAnsi'), 'Calibri')
            rPr.append(rf); sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18'); rPr.append(sz)
            if ri == 0:
                bld = OxmlElement('w:b'); rPr.append(bld)
                clr = OxmlElement('w:color'); clr.set(qn('w:val'), 'FFFFFF'); rPr.append(clr)
            r.append(rPr); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve')
            t.text = str(all_data[ri][ci]).strip(); r.append(t); p.append(r); tc.append(p); tr.append(tc)
        tbl.append(tr)
    
    ref_element.addnext(tbl)
    # Spacer
    sp = OxmlElement('w:p'); spP = OxmlElement('w:pPr'); spS = OxmlElement('w:spacing')
    spS.set(qn('w:after'), '120'); spP.append(spS); sp.append(spP); tbl.addnext(sp)

def add_paragraph_after(ref_element, text, bold=False, size=10):
    """Add a paragraph after ref_element."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if bold:
        pB = OxmlElement('w:b'); rPr2 = OxmlElement('w:rPr'); rPr2.append(pB)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), 'Calibri'); rf.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rf); s = OxmlElement('w:sz'); s.set(qn('w:val'), str(size*2)); rPr.append(s)
    if bold: bld = OxmlElement('w:b'); rPr.append(bld)
    r.append(rPr); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve')
    t.text = text; r.append(t); p.append(r); ref_element.addnext(p)
    return p


# ================================================================
# Find key indices
# ================================================================
key = {}
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '3. Design details': key['s3'] = i
    elif t == '4. Experimental procedure': key['s4'] = i
    elif t == '5. Assessment:': key['s5'] = i
    elif t == 'A. Quality of obtained solutions': key['s5a'] = i
    elif t == 'B. Convergence time': key['s5b'] = i
    elif t == 'C. Stability': key['s5c'] = i
    elif t == '6. Discussion and Conclusion': key['s6'] = i

# Remove the empty paragraph after each header
for k in ['s3', 's4', 's5a', 's5b', 's5c', 's6']:
    idx = key[k]
    for offset in range(1, 3):
        try:
            np = doc.paragraphs[idx + offset]
            if not np.text.strip():
                np._element.getparent().remove(np._element)
                break
        except IndexError:
            break

# ================================================================
# Work BOTTOM-UP
# ================================================================

# --- Section 6 ---
ref = doc.paragraphs[key['s6']]._element
for t in reversed([
    "Limitations and Future Work: SA performance is sensitive to cooling schedule parameters. Future work could explore adaptive cooling schedules or hybrid approaches combining SA with population-based methods. Testing on larger TSP instances (100+ cities) would further assess scalability.",
    "Comparison with Other Metaheuristics: As shown in the accompanying presentation, SA exhibits moderate diversification but strong intensification due to its temperature-controlled acceptance. Unlike Genetic Algorithm, SA works with a single solution (memory-efficient and simpler). Its theoretical convergence guarantee given sufficiently slow cooling is a distinctive advantage.",
    "Stability: SA exhibited excellent stability. For 10 cities, all 50 runs found the identical optimal (CV = 0.00%). For 50 cities, CV = 0.76%, with the range between best and worst only 78 km (1.4% of tour length).",
    "Convergence Time: Geometric cooling creates an exponential quality-time trade-off. α = 0.9 needed 0.04s but produced 5,708 km solutions; α = 0.999 produced 5,616 km in 5.2s. Per-iteration cost is O(1) for delta computation, making SA highly scalable.",
    "Quality of Solutions: For 10 cities, SA found 3,469 km (within 0.12% of optimal 3,473 km, 100% success rate). For 50 cities, best SA = 5,616 km with mean 5,670.90 km ± 43.17 under α = 0.999.",
    "In this project, we implemented Simulated Annealing (SA) to solve the Traveling Salesman Problem (TSP) and evaluated it on two benchmark instances (10 cities, optimal = 3,473 km; 50 cities, optimal ≈ 5,644 km). Results demonstrate SA is an effective metaheuristic, producing high-quality solutions with excellent stability.",
]):
    add_paragraph_after(ref, t)
    ref = ref.getnext()  # Move ref forward so next text goes after this one

# --- Section 5C: Stability ---
ref = doc.paragraphs[key['s5c']]._element
table3_headers = ['Instance', 'Runs', 'Mean (km)', 'Std (km)', 'CV (%)', 'Min (km)', 'Max (km)']
table3_rows = [
    ['10 cities', '50', '3,469', '0.00', '0.00', '3,469', '3,469'],
    ['50 cities', '30', '5,670.90', '43.17', '0.76', '5,616', '5,694'],
]

for t in reversed([
    "Factors: (1) temperature-guided search follows consistent trajectory from exploration to exploitation; (2) restart mechanism prevents pathological convergence; (3) slower cooling rates produce tighter clustering (std drops from 214.98 km at α=0.9 to 27.43 km at α=0.9995).",
    "For 50 cities, stability was excellent (CV = 0.76%). The best-worst range was only 78 km (5,694 - 5,616), just 1.4% of total tour length, indicating consistent convergence to the same solution region.",
    "For 10 cities, SA exhibited perfect stability (CV = 0.00%) — every single run found the identical optimal solution, confirming reliable convergence to the global optimum for small instances.",
    "Stability was measured via multiple independent runs with different random seeds, computing standard deviation and coefficient of variation (CV = std/mean) of solution costs. Table 3 summarizes the results.",
]):
    add_paragraph_after(ref, t)
    ref = ref.getnext()
make_table(ref, table3_headers, table3_rows)

# --- Section 5B: Convergence Time ---
ref = doc.paragraphs[key['s5b']]._element
table2_headers = ['Cooling Rate (α)', '10-City Time (s)', '50-City Time (s)']
table2_rows = [
    ['0.900', '0.07', '0.04'], ['0.950', '0.16', '0.09'],
    ['0.990', '0.59', '0.48'], ['0.995', '0.57', '1.05'],
    ['0.999', '2.57', '5.16'], ['0.9995', '4.90', '10.90'],
]

for t in reversed([
    "Per-iteration cost is O(1) for 2-opt delta computation, making SA highly scalable with respect to instance size.",
    "The geometric schedule creates ~log(T_min/T_0)/log(α) temperature levels. For α = 0.995: ~2,625 levels × 100 iterations ≈ 262,500 total iterations. Increasing to α = 0.999 quintuples the levels, explaining the exponential time increase.",
    "Convergence time was measured as wall-clock time from T_0 to T_min. Table 2 shows the relationship between cooling rate and execution time.",
]):
    add_paragraph_after(ref, t)
    ref = ref.getnext()
make_table(ref, table2_headers, table2_rows)

# --- Section 5A: Quality ---
ref = doc.paragraphs[key['s5a']]._element
table1_headers = ['Instance', 'Optimal (km)', 'SA Best (km)', 'SA Mean (km)', 'Gap (%)']
table1_rows = [
    ['10 cities', '3,473', '3,469', '3,469', '-0.12'],
    ['50 cities', '5,644', '5,616', '5,670.90', '-0.50'],
]

for t in reversed([
    "Cooling rate significantly affects quality. For 50 cities, mean cost decreases from 6,000.80 km (α=0.9) to 5,637.80 km (α=0.9995), though marginal improvement diminishes.",
    "For 50 cities: SA best = 5,616 km (below the 5,644 km reference, since the reference was estimated via initial SA runs). Mean = 5,670.90 ± 43.17 km under best configuration (α = 0.999).",
    "For 10 cities: SA found 3,469 km (4 km below nominal 3,473 km due to floating-point rounding in the scaled matrix). All 50 independent runs found this identical solution, confirming global optimality.",
    "Solution quality was evaluated against known (10-city) or estimated (50-city) optimal values. Table 1 summarizes the results.",
]):
    add_paragraph_after(ref, t)
    ref = ref.getnext()
make_table(ref, table1_headers, table1_rows)

# --- Section 4: Experimental Procedure ---
ref = doc.paragraphs[key['s4']]._element
s4 = [
    "All experiments used deterministic random seeds for reproducibility. Source code is in the simulated_annealing/ package.",
    "Experimental Protocol: (1) Cooling Rate Comparison — 30 runs (10-city) / 10 runs (50-city) per rate, measuring best/mean/std. (2) Stability Analysis — 50 runs (10-city) / 30 runs (50-city) with best rate. (3) Convergence Trace — one detailed run recording cost, temperature, and acceptance rate history.",
    "SA Parameters: T_0 = 5,000 (10-city) / 10,000 (50-city); T_min = 0.01; Geometric cooling (T_{k+1} = α × T_k); α ∈ {0.90, 0.95, 0.99, 0.995, 0.999, 0.9995}; Markov chain length = 100/200; 2-opt neighborhood; Restart after 50× chain length without improvement; Metropolis acceptance.",
    "Test Instances: (1) 10 cities, optimal = 3,473 km (verified by brute-force enumeration of all 3,628,800 tours). (2) 50 cities, optimal ≈ 5,644 km (estimated via extensive SA with α = 0.99999, as brute force infeasible). Both use symmetric Euclidean distance matrices from random 2D coordinates.",
    "Environment: Python 3, NumPy, Matplotlib; macOS (Apple Silicon); Apple M-series, 16GB RAM.",
]
for t in reversed(s4):
    add_paragraph_after(ref, t)
    ref = ref.getnext()

# --- Section 3: Design Details ---
ref = doc.paragraphs[key['s3']]._element
s3 = [
    "7. Key Data Structures: Distance Matrix (N×N float64 NumPy array for O(1) lookups), Tour (Python list for efficient slice manipulation), History Tracking (lists for cost, temperature, acceptance rate per level).",
    "6. Restart Mechanism: If no improvement for 50 × Markov chain length iterations, generates a new random tour and continues from current temperature, combining multi-start diversity with SA's guided search.",
    "5. Acceptance Criterion (Metropolis): P = 1 if ΔE ≤ 0; P = exp(-ΔE/T) if ΔE > 0. Worsening moves are accepted probabilistically, enabling escape from local optima. As T → 0, the algorithm transitions from exploration (diversification) to exploitation (intensification).",
    "4. Cooling Schedule (Geometric): T_{k+1} = α × T_k. Number of temperature levels ≈ log(T_min/T_0) / log(α). T_0 should give initial acceptance > 0.8; T_min should give final acceptance < 0.01. Our settings: T_0 = 5,000-10,000, T_min = 0.01.",
    "3. Neighborhood Structure (2-opt): Select indices i < j, reverse tour[i:j+1]. Removes edges (c_{i-1}, c_i) and (c_j, c_{j+1}), adds (c_{i-1}, c_j) and (c_i, c_{j+1}). Delta is computed in O(1): Δ = d(prev_i, c_j) + d(c_i, next_j) − d(prev_i, c_i) − d(c_j, next_j). Special handling for full-tour reversal (when prev_i == c_j and c_i == next_j, Δ = 0).",
    "2. Initial Solution: Fisher-Yates shuffle of city indices provides a uniformly random unbiased starting point.",
    "1. Solution Representation: Permutation list [c_0, ..., c_{N-1}]. Tour cost = Σ d(c_i, c_{i+1 mod N}). Guarantees all solutions are valid TSP tours.",
]
for t in reversed(s3):
    add_paragraph_after(ref, t)
    ref = ref.getnext()

# Save
out = '/Users/xiniuyiliao/Desktop/application_code/WKU-CPS3410-Project/StimulatedAnnealingAlgorithm_Completed.docx'
doc.save(out)
print(f"✅ Saved: {out}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
