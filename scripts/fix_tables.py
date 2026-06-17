#!/usr/bin/env python3
"""Replace markdown-style text tables with proper Word tables."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc_path = '/Users/xiniuyiliao/Desktop/application_code/WKU-CPS3410-Project/StimulatedAnnealingAlgorithm_Completed.docx'
doc = Document(doc_path)

def insert_table_after(ref_element, headers, rows):
    """Insert a styled Word table right after ref_element in the XML."""
    all_data = [headers] + rows
    num_rows = len(all_data)
    num_cols = len(headers)
    
    tbl = OxmlElement('w:tbl')
    
    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    
    # Borders
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '808080')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tbl.append(tblPr)
    
    # Grid
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(9000 // num_cols))
        tblGrid.append(gc)
    tbl.append(tblGrid)
    
    for ri in range(num_rows):
        tr = OxmlElement('w:tr')
        for ci in range(num_cols):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(9000 // num_cols))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            
            # Cell shading
            shd = OxmlElement('w:shd')
            if ri == 0:
                shd.set(qn('w:fill'), '2F5496')
            elif ri % 2 == 0:
                shd.set(qn('w:fill'), 'F2F2F2')
            else:
                shd.set(qn('w:fill'), 'FFFFFF')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
            tc.append(tcPr)
            
            # Paragraph
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            p.append(pPr)
            
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), 'Calibri')
            rf.set(qn('w:hAnsi'), 'Calibri')
            rPr.append(rf)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')
            rPr.append(sz)
            
            if ri == 0:
                bld = OxmlElement('w:b')
                rPr.append(bld)
                clr = OxmlElement('w:color')
                clr.set(qn('w:val'), 'FFFFFF')
                rPr.append(clr)
            
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = str(all_data[ri][ci]).strip()
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    
    ref_element.addnext(tbl)
    
    # Add spacer after table
    spacer = OxmlElement('w:p')
    spPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:after'), '120')
    spPr.append(sp)
    spacer.append(spPr)
    tbl.addnext(spacer)


# ================================================================
# Find all markdown-style table rows and group them
# ================================================================

# Scan paragraphs to identify:
# - Title: "Table N: ..."
# - Header: has | and contains column names
# - Separator: "---|" pattern
# - Data: has | and numbers
# - Text after: paragraph without |

# We'll create triples of (title_idx, first_data_idx, last_data_idx, header_idx)
table_ranges = []
i = 0
while i < len(doc.paragraphs):
    text = doc.paragraphs[i].text.strip()
    if re.match(r'^Table \d', text):
        title_idx = i
        header_idx = None
        data_start = None
        data_end = None
        
        # Look ahead for header/separator/data rows
        j = i + 1
        while j < len(doc.paragraphs):
            t = doc.paragraphs[j].text.strip()
            if not t:
                j += 1
                continue
            # Check if this is a table row (contains |)
            has_pipe = '|' in t
            is_separator = bool(re.match(r'^-{3,}', t.replace('|', '').replace(' ', '')))
            
            if has_pipe and not is_separator:
                if header_idx is None and data_start is None:
                    # This is the header row
                    header_idx = j
                elif data_start is None:
                    data_start = j
                data_end = j
            elif not has_pipe and data_end is not None:
                # Non-table text after table data → end of table
                break
            elif not has_pipe and not is_separator and data_end is None and j > i + 5:
                # No table found after title within reasonable range
                break
            j += 1
        
        if data_start is not None:
            table_ranges.append((title_idx, header_idx, data_start, data_end))
            i = data_end + 1
        else:
            i += 1
    else:
        i += 1

print(f"Found {len(table_ranges)} table ranges:")
for tr in table_ranges:
    print(f"  Title P{tr[0]}: {doc.paragraphs[tr[0]].text.strip()[:80]}")
    if tr[1] is not None:
        print(f"  Header P{tr[1]}: {doc.paragraphs[tr[1]].text.strip()[:80]}")
    print(f"  Data P{tr[2]}-P{tr[3]}: {doc.paragraphs[tr[2]].text.strip()[:60]}")

# ================================================================
# Expected tables (must match order of appearance)
# ================================================================
expected_tables = [
    {
        'headers': ['Instance', 'Optimal (km)', 'SA Best (km)', 'SA Mean (km)', 'Gap (%)'],
        'rows': [
            ['10 cities', '3,473', '3,469', '3,469', '-0.12'],
            ['50 cities', '5,644', '5,616', '5,670.90', '-0.50'],
        ]
    },
    {
        'headers': ['Cooling Rate (α)', '10-City Time (s)', '50-City Time (s)'],
        'rows': [
            ['0.900', '0.07', '0.04'],
            ['0.950', '0.16', '0.09'],
            ['0.990', '0.59', '0.48'],
            ['0.995', '0.57', '1.05'],
            ['0.999', '2.57', '5.16'],
            ['0.9995', '4.90', '10.90'],
        ]
    },
    {
        'headers': ['Instance', 'Runs', 'Mean (km)', 'Std (km)', 'CV (%)', 'Min (km)', 'Max (km)'],
        'rows': [
            ['10 cities', '50', '3,469', '0.00', '0.00', '3,469', '3,469'],
            ['50 cities', '30', '5,670.90', '43.17', '0.76', '5,616', '5,694'],
        ]
    }
]

# Process from bottom to top to maintain indices
# Reverse both lists to keep correct pairing
table_ranges.sort(key=lambda tr: tr[0], reverse=True)
expected_reversed = list(reversed(expected_tables))

for (title_idx, header_idx, data_start, data_end), expected in zip(table_ranges, expected_reversed):
    title_para = doc.paragraphs[title_idx]
    print(f"\nReplacing table at P{title_idx}: '{title_para.text.strip()[:60]}'")
    
    # Collect all indices to remove: header + separator + data rows
    indices_to_remove = []
    if header_idx is not None:
        # Remove header, separator between header and data, and all data rows
        for idx in range(header_idx, data_end + 1):
            indices_to_remove.append(idx)
    
    print(f"  Removing {len(indices_to_remove)} paragraphs [P{min(indices_to_remove)}-P{max(indices_to_remove)}]")
    
    # Remove from highest index to lowest
    for idx in sorted(indices_to_remove, reverse=True):
        try:
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)
        except Exception as e:
            print(f"  Warning at P{idx}: {e}")
    
    # Insert proper Word table
    insert_table_after(title_para._element, expected['headers'], expected['rows'])
    print(f"  Inserted table: {len(expected['headers'])} cols × {len(expected['rows'])+1} rows")


# Save
output_path = '/Users/xiniuyiliao/Desktop/application_code/WKU-CPS3410-Project/StimulatedAnnealingAlgorithm_Completed.docx'
doc.save(output_path)
print(f"\n✅ Saved: {output_path}")
